from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_CHUNKING_STRATEGIES = frozenset({"policy", "faq", "parent_child", "semantic"})
SUPPORTED_SOURCE_SUFFIXES = frozenset({".txt", ".md", ".pdf", ".docx", ".json"})


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    content: str
    doc_type: str


class DocumentLoader:
    """Load supported office documents and resolve their RAG chunking strategy."""

    def load(self, path: str | Path) -> str:
        file_path = Path(path)
        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_SOURCE_SUFFIXES:
            raise ValueError(f"unsupported document type: {suffix}")
        if suffix in {".txt", ".md", ".json"}:
            return file_path.read_text(encoding="utf-8")
        if suffix == ".pdf":
            return self._load_pdf(file_path)
        if suffix == ".docx":
            return self._load_docx(file_path)
        raise ValueError(f"unsupported document type: {suffix}")

    @staticmethod
    def _load_pdf(path: Path) -> str:
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join(page for page in pages if page)

    @staticmethod
    def _load_docx(path: Path) -> str:
        document = DocxDocument(str(path))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return "\n".join(blocks)

    def resolve_type(self, path: str | Path, corpus_root: str | Path | None = None) -> str:
        """Resolve the intended chunking strategy from the corpus folder."""
        file_path = Path(path)
        if corpus_root is not None:
            root = Path(corpus_root).resolve()
            resolved_path = file_path.resolve()
            try:
                relative = resolved_path.relative_to(root)
            except ValueError:
                return DocumentTypeClassifier().classify(file_path.name, self.load(file_path))
            if len(relative.parts) >= 2:
                strategy = relative.parts[0].lower()
                if strategy not in SUPPORTED_CHUNKING_STRATEGIES:
                    raise ValueError(
                        f"unsupported RAG chunking strategy folder: {strategy!r}; "
                        f"expected one of {sorted(SUPPORTED_CHUNKING_STRATEGIES)}"
                    )
                return strategy
        return DocumentTypeClassifier().classify(file_path.name, self.load(file_path))

    def load_corpus(self, corpus_root: str | Path) -> list[LoadedDocument]:
        """Read all supported documents and resolve their explicit/fallback strategy."""
        root = Path(corpus_root).resolve()
        if not root.is_dir():
            raise ValueError(f"RAG corpus directory does not exist: {root}")
        documents: list[LoadedDocument] = []
        for path in sorted(root.rglob("*")):
            if not path.is_file() or path.suffix.lower() not in SUPPORTED_SOURCE_SUFFIXES:
                continue
            content = self.load(path)
            documents.append(
                LoadedDocument(
                    path=path,
                    content=content,
                    doc_type=self.resolve_type(path, root),
                )
            )
        return documents


class DocumentTypeClassifier:
    def classify(self, name: str, text: str) -> str:
        normalized = text.lower()
        if "第" in text and "条" in text:
            return "policy"
        if name.lower().endswith((".faq", ".faq.md")) or "faq" in normalized[:200]:
            return "faq"
        if "summary" in normalized[:500] or "摘要" in text[:500]:
            return "parent_child"
        return "semantic"
