from pathlib import Path
import json

import pytest
from docx import Document
from pypdf import PdfWriter

from app.rag.loader import DocumentLoader


def test_resolve_type_uses_explicit_strategy_folder(tmp_path: Path) -> None:
    corpus = tmp_path / "rag"
    for strategy in ("policy", "faq", "parent_child", "semantic"):
        document = corpus / strategy / "example.md"
        document.parent.mkdir(parents=True)
        document.write_text("generic content", encoding="utf-8")
        assert DocumentLoader().resolve_type(document, corpus) == strategy


def test_resolve_type_falls_back_to_content_classifier(tmp_path: Path) -> None:
    corpus = tmp_path / "rag"
    corpus.mkdir()
    document = corpus / "rules.md"
    document.write_text("第一条：适用范围", encoding="utf-8")

    assert DocumentLoader().resolve_type(document, corpus) == "policy"


def test_resolve_type_rejects_unknown_strategy_folder(tmp_path: Path) -> None:
    corpus = tmp_path / "rag"
    document = corpus / "unknown" / "example.md"
    document.parent.mkdir(parents=True)
    document.write_text("generic content", encoding="utf-8")

    with pytest.raises(ValueError, match="unsupported RAG chunking strategy folder"):
        DocumentLoader().resolve_type(document, corpus)


def test_resolve_type_works_without_corpus_root(tmp_path: Path) -> None:
    document = tmp_path / "faq.md"
    document.write_text("FAQ: how does this work?", encoding="utf-8")

    assert DocumentLoader().resolve_type(document) == "faq"


def test_load_corpus_reads_all_strategy_folders(tmp_path: Path) -> None:
    corpus = tmp_path / "rag"
    for strategy in ("policy", "faq", "parent_child", "semantic"):
        document = corpus / strategy / f"{strategy}.md"
        document.parent.mkdir(parents=True)
        document.write_text(f"content for {strategy}", encoding="utf-8")

    documents = DocumentLoader().load_corpus(corpus)

    assert {document.doc_type for document in documents} == {
        "policy",
        "faq",
        "parent_child",
        "semantic",
    }
    assert {document.path.parent.name for document in documents} == {
        "policy",
        "faq",
        "parent_child",
        "semantic",
    }


def test_load_pdf_document(tmp_path: Path) -> None:
    pdf_path = tmp_path / "semantic" / "example.pdf"
    pdf_path.parent.mkdir(parents=True)
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with pdf_path.open("wb") as handle:
        writer.write(handle)

    assert DocumentLoader().load(pdf_path) == ""


def test_load_docx_document(tmp_path: Path) -> None:
    docx_path = tmp_path / "parent_child" / "example.docx"
    docx_path.parent.mkdir(parents=True)
    document = Document()
    document.add_paragraph("企业知识库测试内容")
    document.save(docx_path)

    assert DocumentLoader().load(docx_path) == "企业知识库测试内容"


def test_load_corpus_reads_pdf_and_docx_with_folder_strategy(tmp_path: Path) -> None:
    corpus = tmp_path / "rag"

    pdf_path = corpus / "policy" / "policy.pdf"
    pdf_path.parent.mkdir(parents=True)
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with pdf_path.open("wb") as handle:
        writer.write(handle)

    docx_path = corpus / "parent_child" / "parent.docx"
    docx_path.parent.mkdir(parents=True)
    document = Document()
    document.add_paragraph("父子文档测试内容")
    document.save(docx_path)

    documents = DocumentLoader().load_corpus(corpus)

    assert {(document.path.suffix, document.doc_type) for document in documents} == {
        (".pdf", "policy"),
        (".docx", "parent_child"),
    }
    assert next(document for document in documents if document.path.suffix == ".docx").content == "父子文档测试内容"


def test_load_policy_json_array(tmp_path: Path) -> None:
    path = tmp_path / "policy" / "labor.json"
    path.parent.mkdir(parents=True)
    payload = [
        {"中华人民共和国劳动合同法 第一条": "完善劳动合同制度。"},
        {"中华人民共和国劳动合同法 第二条": "保护劳动者合法权益。"},
    ]
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    assert DocumentLoader().load(path) == json.dumps(payload, ensure_ascii=False)


def test_policy_nodes_parse_json_array_into_article_chunks(tmp_path: Path) -> None:
    from app.rag.splitter import policy_nodes

    path = tmp_path / "policy" / "labor.json"
    path.parent.mkdir(parents=True)
    payload = [
        {"中华人民共和国劳动合同法 第一条": "完善劳动合同制度。"},
        {"中华人民共和国劳动合同法 第二条": "保护劳动者合法权益。"},
    ]
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    chunks = policy_nodes(DocumentLoader().load(path), path.name)

    assert len(chunks) == 2
    assert chunks[0].content == "中华人民共和国劳动合同法 第一条：完善劳动合同制度。"
    assert chunks[0].metadata["document"] == "中华人民共和国劳动合同法"
    assert chunks[0].metadata["article"] == "第一条"
    assert chunks[1].metadata["article"] == "第二条"
